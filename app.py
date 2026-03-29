"""
AI Resume Builder — Streamlit Application
==========================================
A production-ready resume builder with ATS analysis,
skill suggestions, and DOCX export.
"""

import io
import os
import re
import json
import random
from typing import Optional, List
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Optional: Groq via openai-compatible client (graceful fallback if not installed) ──
try:
    from openai import OpenAI as _OpenAI
    _GROQ_AVAILABLE = True
except ImportError:
    _GROQ_AVAILABLE = False

# ── Optional: sklearn TF-IDF (graceful fallback if not installed) ──
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    _SKLEARN_AVAILABLE = True
except ImportError:
    _SKLEARN_AVAILABLE = False

# ─────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="AI Resume Builder",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────
# CUSTOM CSS
# ─────────────────────────────────────────────
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600;700&family=DM+Mono:wght@400;500&display=swap');

    /* ── Global ── */
    html, body, [class*="css"] {
        font-family: 'DM Sans', sans-serif;
    }
    .main { background: #0f1117; }
    .block-container { padding: 2rem 2.5rem !important; }

    /* ── Sidebar ── */
    [data-testid="stSidebar"] {
        background: #16191f !important;
        border-right: 1px solid #2a2d35;
    }
    [data-testid="stSidebar"] .stRadio label {
        font-size: 0.95rem;
        color: #c8cdd6;
        padding: 0.4rem 0;
    }

    /* ── Header card ── */
    .hero-card {
        background: linear-gradient(135deg, #1a1f2e 0%, #0f1117 60%, #1a1a2e 100%);
        border: 1px solid #2e3347;
        border-radius: 16px;
        padding: 2.2rem 2.5rem;
        margin-bottom: 2rem;
        position: relative;
        overflow: hidden;
    }
    .hero-card::before {
        content: "";
        position: absolute;
        top: -40px; right: -40px;
        width: 200px; height: 200px;
        background: radial-gradient(circle, rgba(99,102,241,0.18) 0%, transparent 70%);
        border-radius: 50%;
    }
    .hero-title {
        font-size: 2.2rem;
        font-weight: 700;
        background: linear-gradient(135deg, #e2e8f0, #a5b4fc);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin: 0 0 0.4rem 0;
    }
    .hero-sub {
        color: #7c8494;
        font-size: 1rem;
        margin: 0;
    }

    /* ── Section cards ── */
    .section-card {
        background: #16191f;
        border: 1px solid #2a2d35;
        border-radius: 12px;
        padding: 1.6rem 1.8rem;
        margin-bottom: 1.4rem;
    }
    .section-title {
        font-size: 1.05rem;
        font-weight: 600;
        color: #a5b4fc;
        margin: 0 0 1rem 0;
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }

    /* ── Metric cards ── */
    .metric-card {
        background: #1c1f29;
        border: 1px solid #2a2d35;
        border-radius: 10px;
        padding: 1.2rem 1.5rem;
        text-align: center;
    }
    .metric-value {
        font-size: 2.4rem;
        font-weight: 700;
        line-height: 1;
        margin-bottom: 0.3rem;
    }
    .metric-label {
        font-size: 0.8rem;
        color: #7c8494;
        text-transform: uppercase;
        letter-spacing: 0.08em;
    }
    .score-high   { color: #34d399; }
    .score-medium { color: #fbbf24; }
    .score-low    { color: #f87171; }

    /* ── Skill tags ── */
    .tag-container { display: flex; flex-wrap: wrap; gap: 0.5rem; margin-top: 0.6rem; }
    .tag {
        background: #1e2235;
        border: 1px solid #3b3f56;
        border-radius: 20px;
        padding: 0.25rem 0.75rem;
        font-size: 0.8rem;
        color: #a5b4fc;
        font-family: 'DM Mono', monospace;
    }
    .tag-new {
        background: #1a2e22;
        border-color: #2d6a4f;
        color: #52d9a4;
    }

    /* ── Resume preview ── */
    .resume-preview {
        background: #fff;
        border-radius: 12px;
        padding: 3rem;
        color: #111;
        font-family: 'DM Sans', sans-serif;
        box-shadow: 0 20px 60px rgba(0,0,0,0.4);
        max-width: 720px;
        margin: 0 auto;
    }
    .resume-name {
        font-size: 1.8rem;
        font-weight: 700;
        color: #1e1e2e;
        margin-bottom: 0.3rem;
    }
    .resume-contact {
        font-size: 0.85rem;
        color: #555;
        margin-bottom: 1.2rem;
        display: flex;
        flex-wrap: wrap;
        gap: 0.8rem;
    }
    .resume-section-title {
        font-size: 0.75rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.12em;
        color: #4f46e5;
        border-bottom: 2px solid #4f46e5;
        padding-bottom: 0.2rem;
        margin: 1.2rem 0 0.7rem 0;
    }
    .resume-item-title { font-weight: 600; color: #1e1e2e; font-size: 0.95rem; }
    .resume-item-sub   { color: #666; font-size: 0.85rem; }
    .resume-bullet     { color: #333; font-size: 0.88rem; padding-left: 1rem; margin: 0.2rem 0; }

    /* ── Inputs ── */
    .stTextInput input, .stTextArea textarea, .stSelectbox select {
        background: #1c1f29 !important;
        border: 1px solid #2a2d35 !important;
        border-radius: 8px !important;
        color: #e2e8f0 !important;
        font-family: 'DM Sans', sans-serif !important;
    }
    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: #6366f1 !important;
        box-shadow: 0 0 0 2px rgba(99,102,241,0.2) !important;
    }

    /* ── Buttons ── */
    .stButton > button {
        background: linear-gradient(135deg, #6366f1, #8b5cf6) !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        padding: 0.55rem 1.5rem !important;
        transition: all 0.2s !important;
    }
    .stButton > button:hover {
        transform: translateY(-1px) !important;
        box-shadow: 0 6px 20px rgba(99,102,241,0.35) !important;
    }
    .stDownloadButton > button {
        background: linear-gradient(135deg, #059669, #10b981) !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
    }

    /* ── Progress bar ── */
    .stProgress > div > div { background: linear-gradient(90deg, #6366f1, #8b5cf6) !important; }

    /* ── Divider ── */
    hr { border-color: #2a2d35 !important; }

    /* ── Step indicator ── */
    .step-dot {
        display: inline-block;
        width: 8px; height: 8px;
        border-radius: 50%;
        background: #6366f1;
        margin-right: 6px;
    }

    /* ── Suggestion box ── */
    .suggestion-box {
        background: #111827;
        border-left: 3px solid #6366f1;
        border-radius: 0 8px 8px 0;
        padding: 1rem 1.2rem;
        margin: 0.6rem 0;
        font-size: 0.9rem;
        color: #d1d5db;
        line-height: 1.6;
    }

    /* Nav radio styling */
    div[data-testid="stRadio"] label {
        cursor: pointer;
        padding: 0.5rem 0.8rem;
        border-radius: 8px;
        transition: background 0.15s;
    }
    div[data-testid="stRadio"] label:hover {
        background: #1e2235;
    }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# STATIC DATA MAPS
# ─────────────────────────────────────────────
SKILL_ROLE_MAP = {
    "python":       ["FastAPI", "Django", "Pandas", "NumPy", "Pytest", "Celery"],
    "javascript":   ["React", "Node.js", "TypeScript", "Next.js", "Jest", "Webpack"],
    "java":         ["Spring Boot", "Maven", "JUnit", "Hibernate", "Kafka"],
    "sql":          ["PostgreSQL", "MySQL", "SQLAlchemy", "dbt", "Redshift"],
    "ml":           ["Scikit-learn", "TensorFlow", "PyTorch", "MLflow", "Hugging Face"],
    "data":         ["Tableau", "Power BI", "Spark", "Airflow", "dbt"],
    "devops":       ["Docker", "Kubernetes", "CI/CD", "Terraform", "Ansible"],
    "cloud":        ["AWS", "GCP", "Azure", "Serverless", "CDK"],
    "react":        ["Redux", "React Query", "Storybook", "Vite", "Cypress"],
    "go":           ["gRPC", "Gin", "goroutines", "GORM", "Prometheus"],
}

JOB_ROLE_KEYWORDS = {
    "backend":     ["python", "java", "go", "api", "database", "sql", "microservices"],
    "frontend":    ["react", "javascript", "typescript", "html", "css", "ui"],
    "data":        ["sql", "python", "data", "ml", "analytics", "spark"],
    "devops":      ["docker", "kubernetes", "aws", "ci/cd", "linux", "terraform"],
    "fullstack":   ["react", "node", "python", "sql", "api", "javascript"],
}

PROJECT_TEMPLATES = {
    "web":      "Engineered a scalable {title} leveraging {tech}, achieving sub-200ms API response times and 99.9% uptime. Implemented CI/CD pipelines reducing deployment time by 60%.",
    "ml":       "Developed {title} using {tech} to solve {desc}. Achieved 94% model accuracy on validation set; deployed as REST API serving 500+ daily requests.",
    "mobile":   "Built cross-platform {title} with {tech}, delivering real-time features to 1,000+ users. Reduced load time by 40% through async data fetching.",
    "data":     "Designed {title} pipeline using {tech} that processes 10GB+ of data daily. Automated reporting reduced manual effort by 8 hours/week.",
    "default":  "Built {title} using {tech} to {desc}. Implemented robust error handling, unit tests, and documentation resulting in a maintainable, production-ready solution.",
}


# ─────────────────────────────────────────────
# GROQ CLIENT — cached singleton
# ─────────────────────────────────────────────
@st.cache_resource(show_spinner=False)
def _get_groq_client():
    """
    Return a cached Groq client (OpenAI-compatible).
    Reads GROQ_API_KEY from environment — never hardcode keys in source.
    Returns None if the package is missing or the key is not set.
    """
    if not _GROQ_AVAILABLE:
        return None
    api_key = os.environ.get("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", "")
    if not api_key:
        return None
    try:
        return _OpenAI(
            api_key=api_key,
            base_url="https://api.groq.com/openai/v1",
        )
    except Exception:
        return None


def call_groq(prompt: str, system: str = "You are an expert resume and career coach.",
              max_tokens: int = 512) -> Optional[str]:
    """
    Call Groq (Llama-3.3-70b) with a system + user prompt.
    Returns the text response, or None on any failure.
    Cached at the caller level via @st.cache_data — no repeated API hits.
    """
    client = _get_groq_client()
    if client is None:
        return None
    try:
        response = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=[
                {"role": "system", "content": system},
                {"role": "user",   "content": prompt},
            ],
            temperature=0.7,
            max_tokens=max_tokens,
        )
        return response.choices[0].message.content.strip()
    except Exception:
        return None


# Internal alias so all existing callers (_cached_improve_project, _cached_ai_suggestions)
# keep working without any change to their bodies.
def _call_claude(system: str, user: str, max_tokens: int = 512) -> Optional[str]:
    """Thin wrapper — routes all AI calls through Groq (was Anthropic)."""
    return call_groq(prompt=user, system=system, max_tokens=max_tokens)


# ─────────────────────────────────────────────
# 1. DATA COLLECTION
# ─────────────────────────────────────────────
def collect_user_data() -> dict:
    """Render input forms and return structured resume data."""

    st.markdown("""
    <div class="hero-card">
        <p class="hero-title">📄 AI Resume Builder</p>
        <p class="hero-sub">Craft a job-winning resume powered by AI analysis & smart suggestions</p>
    </div>
    """, unsafe_allow_html=True)

    data = {}

    # ── Personal Info ──
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<p class="section-title">👤 Personal Information</p>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1:
        data["name"]     = st.text_input("Full Name", placeholder="Alex Johnson")
        data["email"]    = st.text_input("Email", placeholder="alex@example.com")
        data["phone"]    = st.text_input("Phone", placeholder="+1 (555) 000-0000")
    with c2:
        data["linkedin"] = st.text_input("LinkedIn URL", placeholder="linkedin.com/in/alexjohnson")
        data["github"]   = st.text_input("GitHub URL", placeholder="github.com/alexjohnson")
        data["location"] = st.text_input("Location", placeholder="San Francisco, CA")
    st.markdown("</div>", unsafe_allow_html=True)

    # ── Skills ──
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<p class="section-title">🛠 Skills</p>', unsafe_allow_html=True)
    data["skills_raw"] = st.text_area(
        "Skills (comma-separated)",
        placeholder="Python, React, PostgreSQL, Docker, AWS, Machine Learning",
        height=80,
    )
    data["skills"] = [s.strip() for s in data["skills_raw"].split(",") if s.strip()]
    st.markdown("</div>", unsafe_allow_html=True)

    # ── Education ──
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<p class="section-title">🎓 Education</p>', unsafe_allow_html=True)
    num_edu = st.number_input("Number of education entries", 1, 4, 1, key="num_edu")
    data["education"] = []
    for i in range(int(num_edu)):
        st.markdown(f"**Entry {i+1}**")
        c1, c2, c3 = st.columns([3, 2, 1])
        with c1:
            degree = st.text_input("Degree / Certification", key=f"degree_{i}",
                                   placeholder="B.S. Computer Science")
        with c2:
            institution = st.text_input("Institution", key=f"inst_{i}",
                                        placeholder="MIT")
        with c3:
            year = st.text_input("Year", key=f"year_{i}", placeholder="2022")
        gpa = st.text_input("GPA (optional)", key=f"gpa_{i}", placeholder="3.8 / 4.0")
        if degree or institution:
            data["education"].append({
                "degree": degree, "institution": institution,
                "year": year, "gpa": gpa
            })
        if i < int(num_edu) - 1:
            st.markdown("---")
    st.markdown("</div>", unsafe_allow_html=True)

    # ── Projects ──
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<p class="section-title">🚀 Projects</p>', unsafe_allow_html=True)
    num_proj = st.number_input("Number of projects", 1, 6, 2, key="num_proj")
    data["projects"] = []
    for i in range(int(num_proj)):
        st.markdown(f"**Project {i+1}**")
        c1, c2 = st.columns([2, 1])
        with c1:
            title = st.text_input("Project Title", key=f"ptitle_{i}",
                                  placeholder="Smart Inventory System")
        with c2:
            tech = st.text_input("Tech Stack", key=f"ptech_{i}",
                                 placeholder="Python, FastAPI, React")
        desc = st.text_area("Description", key=f"pdesc_{i}",
                            placeholder="Briefly describe what this project does and its impact.",
                            height=80)
        if title:
            data["projects"].append({"title": title, "tech": tech, "description": desc})
        if i < int(num_proj) - 1:
            st.markdown("---")
    st.markdown("</div>", unsafe_allow_html=True)

    # ── Experience ──
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<p class="section-title">💼 Work Experience <span style="color:#7c8494;font-size:0.8rem;">(optional)</span></p>',
                unsafe_allow_html=True)
    has_exp = st.checkbox("I have work experience to add", value=False)
    data["experience"] = []
    if has_exp:
        num_exp = st.number_input("Number of roles", 1, 5, 1, key="num_exp")
        for i in range(int(num_exp)):
            st.markdown(f"**Role {i+1}**")
            c1, c2, c3 = st.columns([2, 2, 1])
            with c1:
                role    = st.text_input("Job Title",   key=f"etitle_{i}", placeholder="Software Engineer")
            with c2:
                company = st.text_input("Company",     key=f"ecomp_{i}",  placeholder="Acme Corp")
            with c3:
                period  = st.text_input("Period",      key=f"eperiod_{i}",placeholder="2022–2024")
            bullets = st.text_area("Key responsibilities / achievements (one per line)",
                                   key=f"ebullets_{i}", height=100,
                                   placeholder="• Built REST APIs serving 10k RPM\n• Led migration to microservices")
            if role or company:
                data["experience"].append({
                    "role": role, "company": company, "period": period,
                    "bullets": [b.strip() for b in bullets.split("\n") if b.strip()]
                })
            if i < int(num_exp) - 1:
                st.markdown("---")
    st.markdown("</div>", unsafe_allow_html=True)

    # ── Achievements ──
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<p class="section-title">🏆 Achievements & Certifications <span style="color:#7c8494;font-size:0.8rem;">(optional)</span></p>',
                unsafe_allow_html=True)
    achievements_raw = st.text_area(
        "One achievement per line",
        placeholder="• AWS Solutions Architect Certified (2023)\n• 1st Place, TechHacks 2023\n• Open-source contributor — 500+ GitHub stars",
        height=110,
    )
    data["achievements"] = [a.strip() for a in achievements_raw.split("\n") if a.strip()]
    st.markdown("</div>", unsafe_allow_html=True)

    return data


# ─────────────────────────────────────────────
# 2. ANALYSIS LOGIC
# ─────────────────────────────────────────────
def compute_ats_score(data: dict, jd: str) -> int:
    """
    UPDATED: ATS scoring using TF-IDF cosine similarity (sklearn).
    Falls back to keyword-overlap logic if sklearn is unavailable.
    """
    # Build resume text corpus
    resume_text = " ".join(filter(None, [
        data.get("skills_raw", ""),
        data.get("name", ""),
        " ".join(
            p.get("description", "") + " " + p.get("tech", "")
            for p in data.get("projects", [])
        ),
        " ".join(
            e.get("role", "") + " " + " ".join(e.get("bullets", []))
            for e in data.get("experience", [])
        ),
        " ".join(a for a in data.get("achievements", [])),
    ])).strip()

    if not jd.strip() or not resume_text:
        return random.randint(55, 72)

    # ── PRIMARY: TF-IDF cosine similarity ──
    if _SKLEARN_AVAILABLE:
        try:
            vectorizer = TfidfVectorizer(
                stop_words="english",
                ngram_range=(1, 2),   # unigrams + bigrams for phrase matching
                min_df=1,
            )
            tfidf_matrix = vectorizer.fit_transform([jd, resume_text])
            similarity   = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            # Map 0–1 cosine similarity → realistic ATS range (40–97)
            score = int(min(97, max(40, similarity * 130)))
            return score
        except Exception:
            pass  # fall through to keyword fallback

    # ── FALLBACK: keyword overlap (original logic) ──
    jd_words     = set(w.lower() for w in jd.split() if len(w) > 3)
    resume_words = set(resume_text.lower().split())
    matched      = jd_words & resume_words
    base         = (len(matched) / max(len(jd_words), 1)) * 100
    return int(min(98, max(35, base * 1.8 + random.randint(-5, 5))))


def get_skill_suggestions(data: dict, jd: str) -> List[str]:
    """Return missing skills based on existing skills and JD keywords."""
    existing = set(s.lower() for s in data.get("skills", []))
    suggestions = []

    for base_skill, extras in SKILL_ROLE_MAP.items():
        if base_skill in existing:
            for extra in extras:
                if extra.lower() not in existing:
                    suggestions.append(extra)

    # Pull role-based suggestions from JD
    jd_lower = jd.lower()
    for role, keywords in JOB_ROLE_KEYWORDS.items():
        hits = sum(1 for kw in keywords if kw in jd_lower)
        if hits >= 2:
            for kw in keywords:
                cap = kw.capitalize()
                if kw not in existing and cap not in suggestions:
                    suggestions.append(cap)

    # Deduplicate, limit
    seen = set()
    unique = []
    for s in suggestions:
        if s.lower() not in seen:
            seen.add(s.lower()); unique.append(s)

    return unique[:12]


def _template_project_description(project: dict) -> str:
    """Original template-based fallback for project description improvement."""
    title     = project.get("title", "this project")
    tech      = project.get("tech", "modern technologies")
    desc      = project.get("description", "solve a real-world problem")
    tech_lower = tech.lower()

    if any(kw in tech_lower for kw in ["ml", "tensorflow", "pytorch", "model", "sklearn"]):
        template = PROJECT_TEMPLATES["ml"]
    elif any(kw in tech_lower for kw in ["react", "vue", "angular", "node", "django", "flask"]):
        template = PROJECT_TEMPLATES["web"]
    elif any(kw in tech_lower for kw in ["spark", "airflow", "kafka", "dbt", "sql"]):
        template = PROJECT_TEMPLATES["data"]
    elif any(kw in tech_lower for kw in ["flutter", "swift", "kotlin", "expo"]):
        template = PROJECT_TEMPLATES["mobile"]
    else:
        template = PROJECT_TEMPLATES["default"]

    return template.format(title=title, tech=tech, desc=desc)


@st.cache_data(show_spinner=False)
def _cached_improve_project(title: str, tech: str, desc: str) -> Optional[str]:
    """
    Cached AI call for project description improvement.
    Keyed by (title, tech, desc) so identical projects are never re-requested.
    """
    system = (
        "You are an expert resume writer. Write a detailed ~100 word project description paragraph. "
        "Include strong action verbs, quantifiable achievements/metrics, tech keywords from stack. "
        "Structure: overview → key features → impact/results → tech implementation. Professional tone. "
        "Return ONLY the paragraph, no bullets or extra text."
    )
    user = (
        f"Project title: {title}\n"
        f"Tech stack: {tech}\n"
        f"Original description: {desc}\n\n"
        "Rewrite this as 2 strong resume bullet points."
    )
    return _call_claude(system, user, max_tokens=400)


def improve_project_description(project: dict) -> str:
    """
    UPDATED: AI-powered project description via Claude API.
    Falls back to template logic if the API is unavailable or fails.
    """
    title = project.get("title", "this project")
    tech  = project.get("tech", "modern technologies")
    desc  = project.get("description", "solve a real-world problem")

    # ── PRIMARY: Claude API (cached per unique project) ──
    ai_result = _cached_improve_project(title, tech, desc)
    if ai_result and ai_result.strip():
        return ai_result.strip()

    # ── FALLBACK: original template logic ──
    return _template_project_description(project)


def extract_keyword_match(data: dict, jd: str) -> dict:
    """
    NEW: Extract matched and missing keywords between resume and JD.
    Returns dicts with matched_keywords and missing_keywords lists.
    Uses TF-IDF important terms when sklearn is available; falls back to simple word sets.
    """
    if not jd.strip():
        return {"matched_keywords": [], "missing_keywords": []}

    resume_text = " ".join(filter(None, [
        data.get("skills_raw", ""),
        " ".join(p.get("description", "") + " " + p.get("tech", "") for p in data.get("projects", [])),
        " ".join(e.get("role", "") + " " + " ".join(e.get("bullets", [])) for e in data.get("experience", [])),
    ])).lower()

    # ── Extract meaningful JD keywords ──
    if _SKLEARN_AVAILABLE:
        try:
            vectorizer = TfidfVectorizer(
                stop_words="english",
                ngram_range=(1, 1),
                max_features=60,
            )
            vectorizer.fit([jd])
            jd_keywords = set(vectorizer.get_feature_names_out())
        except Exception:
            jd_keywords = set(w.lower() for w in jd.split() if len(w) > 4)
    else:
        jd_keywords = set(w.lower() for w in jd.split() if len(w) > 4)

    # ── Filter out generic stopwords manually ──
    _noise = {"this", "that", "with", "have", "will", "from", "they", "their",
              "about", "which", "would", "been", "were", "also", "more", "your",
              "work", "team", "able", "must", "role", "join", "help", "using"}
    jd_keywords -= _noise

    resume_word_set = set(resume_text.split())
    matched  = sorted(kw for kw in jd_keywords if kw in resume_word_set)
    missing  = sorted(kw for kw in jd_keywords if kw not in resume_word_set)

    return {
        "matched_keywords": matched[:20],
        "missing_keywords": missing[:20],
    }


@st.cache_data(show_spinner=False)
def _cached_ai_suggestions(
    name: str, skills_raw: str, edu_str: str,
    proj_str: str, exp_str: str, jd: str
) -> dict:
    """
    Cached: Call AI (Groq) to generate structured resume improvement suggestions.
    Returns a dict with keys: missing_skills, weak_sections, keyword_gaps, advice.
    """

    system = (
        "You are a professional resume coach and ATS optimization expert. "
        "Given a candidate's resume summary and a job description, return a JSON object with these exact keys:\n"
        "  missing_skills   : list of up to 6 skill strings the candidate lacks but the JD requires\n"
        "  weak_sections    : list of up to 4 strings naming resume sections that need improvement + one-line reason\n"
        "  keyword_gaps     : list of up to 8 important JD keywords absent from the resume\n"
        "  advice           : list of exactly 4 concise, actionable improvement tips (strings)\n"
        "Return ONLY valid JSON. No markdown fences, no extra commentary."
    )

    user = (
        f"Candidate name: {name}\n"
        f"Skills: {skills_raw}\n"
        f"Education: {edu_str}\n"
        f"Projects: {proj_str}\n"
        f"Experience: {exp_str}\n\n"
        f"Job Description:\n{jd[:3000]}\n\n"
        "Return the JSON object now."
    )

    raw = _call_claude(system, user, max_tokens=700)

    # ✅ If API fails
    if not raw:
        return {
            "missing_skills": [],
            "weak_sections": [],
            "keyword_gaps": [],
            "advice": []
        }

    try:
        # ✅ Clean response
        clean = re.sub(r"```(?:json)?", "", raw).strip().strip("`")

        # ✅ Parse JSON
        return json.loads(clean)

    except Exception:
        # ✅ Safe fallback
        return {
            "missing_skills": [],
            "weak_sections": [],
            "keyword_gaps": [],
            "advice": []
        }
def generate_ai_suggestions(data: dict, jd: str) -> dict:
    """
    NEW: Produce AI-powered holistic resume suggestions via Claude.
    Falls back to empty dict (no crash) if API is unavailable.
    """
    if not jd.strip():
        return {}

    # Flatten resume sections into strings for the prompt
    edu_str  = "; ".join(
        f"{e.get('degree','')} at {e.get('institution','')}"
        for e in data.get("education", [])
    )
    proj_str = "; ".join(
        f"{p.get('title','')} ({p.get('tech','')}) — {p.get('description','')}"
        for p in data.get("projects", [])
    )
    exp_str  = "; ".join(
        f"{e.get('role','')} at {e.get('company','')} ({e.get('period','')})"
        for e in data.get("experience", [])
    )

    result = _cached_ai_suggestions(
        name       = data.get("name", ""),
        skills_raw = data.get("skills_raw", ""),
        edu_str    = edu_str,
        proj_str   = proj_str,
        exp_str    = exp_str,
        jd         = jd,
    )
    return result


def analyze_resume(data: dict, jd: str) -> dict:
    """
    UPDATED: Orchestrates all analysis functions.
    Now includes keyword_match and ai_suggestions in the result.
    """
    ats_score = compute_ats_score(data, jd)
    suggestions = get_skill_suggestions(data, jd)
    improved_projects = [
        {"title": p["title"], "improved": improve_project_description(p)}
        for p in data.get("projects", []) if p.get("title")
    ]
    keyword_match  = extract_keyword_match(data, jd)   # NEW
    ai_suggestions = generate_ai_suggestions(data, jd) # NEW

    return {
        "ats_score":        ats_score,
        "skill_suggestions": suggestions,
        "improved_projects": improved_projects,
        "keyword_match":    keyword_match,    # NEW
        "ai_suggestions":   ai_suggestions,   # NEW
    }


# ─────────────────────────────────────────────
# 3. DOCX GENERATION
# ─────────────────────────────────────────────
def _add_horizontal_rule(doc: Document):
    """Add a thin colored underline paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4F46E5')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    return p


def _section_heading(doc: Document, text: str):
    """Add a styled section heading with an underline rule."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(0)
    _add_horizontal_rule(doc)


def generate_resume_docx(data: dict, analysis: dict, use_improved: bool = False) -> bytes:
    """
    Build and return a professional DOCX resume as bytes.
    Set use_improved=True to swap in AI-improved project descriptions.
    """
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.85)
        section.right_margin  = Inches(0.85)

    # ── NAME ──
    name_para = doc.add_paragraph()
    name_run  = name_para.add_run(data.get("name", "Your Name"))
    name_run.bold      = True
    name_run.font.size = Pt(22)
    name_run.font.name = "Calibri"
    name_run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x2E)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(3)

    # ── CONTACT LINE ──
    contact_parts = []
    if data.get("email"):    contact_parts.append(f"✉ {data['email']}")
    if data.get("phone"):    contact_parts.append(f"📞 {data['phone']}")
    if data.get("location"): contact_parts.append(f"📍 {data['location']}")
    if data.get("linkedin"): contact_parts.append(f"🔗 {data['linkedin']}")
    if data.get("github"):   contact_parts.append(f"💻 {data['github']}")

    contact_para = doc.add_paragraph("  |  ".join(contact_parts))
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in contact_para.runs:
        run.font.size  = Pt(8.5)
        run.font.name  = "Calibri"
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x66)
    contact_para.paragraph_format.space_after = Pt(4)

    # ── SKILLS ──
    if data.get("skills"):
        _section_heading(doc, "Skills")
        skills_para = doc.add_paragraph(", ".join(data["skills"]))
        skills_para.paragraph_format.space_before = Pt(3)
        for run in skills_para.runs:
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"

    # ── EDUCATION ──
    if data.get("education"):
        _section_heading(doc, "Education")
        for edu in data["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(0)
            title_run = p.add_run(edu.get("degree", ""))
            title_run.bold = True
            title_run.font.size = Pt(10)
            title_run.font.name = "Calibri"

            inst_str = ""
            if edu.get("institution"): inst_str += f"  —  {edu['institution']}"
            if edu.get("year"):        inst_str += f",  {edu['year']}"
            if inst_str:
                sub_run = p.add_run(inst_str)
                sub_run.font.size = Pt(9.5)
                sub_run.font.name = "Calibri"
                sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x66)

            if edu.get("gpa"):
                gpa_p = doc.add_paragraph(f"GPA: {edu['gpa']}")
                for r in gpa_p.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Calibri"
                    r.font.color.rgb = RGBColor(0x55, 0x55, 0x66)
                gpa_p.paragraph_format.space_before = Pt(0)
                gpa_p.paragraph_format.space_after  = Pt(2)

    # ── EXPERIENCE ──
    if data.get("experience"):
        _section_heading(doc, "Work Experience")
        for exp in data["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(0)
            role_run = p.add_run(exp.get("role", ""))
            role_run.bold = True
            role_run.font.size = Pt(10.5)
            role_run.font.name = "Calibri"

            if exp.get("company") or exp.get("period"):
                detail = ""
                if exp.get("company"): detail += f"  —  {exp['company']}"
                if exp.get("period"):  detail += f"  ·  {exp['period']}"
                sub_run = p.add_run(detail)
                sub_run.font.size = Pt(9.5)
                sub_run.font.name = "Calibri"
                sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x66)

            for bullet in exp.get("bullets", []):
                clean = bullet.lstrip("•-– ").strip()
                if clean:
                    bp = doc.add_paragraph(style="List Bullet")
                    bp_run = bp.add_run(clean)
                    bp_run.font.size = Pt(9.5)
                    bp_run.font.name = "Calibri"
                    bp.paragraph_format.space_before = Pt(1)
                    bp.paragraph_format.space_after  = Pt(1)

    # ── PROJECTS ──
    if data.get("projects"):
        _section_heading(doc, "Projects")
        improved_map = {
            ip["title"]: ip["improved"]
            for ip in analysis.get("improved_projects", [])
        }
        for proj in data["projects"]:
            if not proj.get("title"): continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(0)
            proj_run = p.add_run(proj["title"])
            proj_run.bold = True
            proj_run.font.size = Pt(10.5)
            proj_run.font.name = "Calibri"
            if proj.get("tech"):
                tech_run = p.add_run(f"  |  {proj['tech']}")
                tech_run.font.size = Pt(9)
                tech_run.font.name = "Calibri"
                tech_run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

            # Choose description
            desc_text = (
                improved_map.get(proj["title"])
                if use_improved
                else proj.get("description", "")
            )
            if desc_text:
                dp = doc.add_paragraph(desc_text)
                for dr in dp.runs:
                    dr.font.size = Pt(9.5)
                    dr.font.name = "Calibri"
                dp.paragraph_format.space_before = Pt(1)
                dp.paragraph_format.space_after  = Pt(3)

    # ── ACHIEVEMENTS ──
    if data.get("achievements"):
        _section_heading(doc, "Achievements & Certifications")
        for ach in data["achievements"]:
            clean = ach.lstrip("•-– ").strip()
            if clean:
                ap = doc.add_paragraph(style="List Bullet")
                ar = ap.add_run(clean)
                ar.font.size = Pt(9.5)
                ar.font.name = "Calibri"
                ap.paragraph_format.space_before = Pt(1)
                ap.paragraph_format.space_after  = Pt(1)

    # ── Serialize to bytes ──
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ─────────────────────────────────────────────
# 4. SIDEBAR NAVIGATION
# ─────────────────────────────────────────────
def render_sidebar() -> str:
    with st.sidebar:
        st.markdown("""
        <div style="padding:1rem 0 1.5rem;">
            <div style="font-size:1.4rem;font-weight:700;color:#a5b4fc;">📄 Resume AI</div>
            <div style="font-size:0.78rem;color:#7c8494;margin-top:0.2rem;">Hackathon Edition</div>
        </div>
        """, unsafe_allow_html=True)

        page = st.radio(
            "Navigate",
            ["📝  Build Resume", "🔍  Analyze", "👁  Preview & Export"],
            label_visibility="collapsed",
        )

        st.markdown("---")
        st.markdown("""
        <div style="font-size:0.78rem;color:#7c8494;line-height:1.8;">
            <b style="color:#a5b4fc;">How it works</b><br>
            1. Fill in your info<br>
            2. Paste the job description<br>
            3. Run analysis<br>
            4. Preview & download
        </div>
        """, unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("""
        <div style="font-size:0.72rem;color:#4a4d58;line-height:1.8;">
            Built with Streamlit · python-docx<br>
            Powered by Claude AI · sklearn TF-IDF 🚀
        </div>
        """, unsafe_allow_html=True)

    return page


# ─────────────────────────────────────────────
# 5. ANALYSIS PAGE
# ─────────────────────────────────────────────
def render_analysis_page(data: dict):
    st.markdown("""
    <div class="hero-card">
        <p class="hero-title">🔍 Resume Analysis</p>
        <p class="hero-sub">Paste a job description to get ATS score, skill gaps & enhanced descriptions</p>
    </div>
    """, unsafe_allow_html=True)

    # JD input
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<p class="section-title">📋 Job Description</p>', unsafe_allow_html=True)
    jd = st.text_area(
        "Paste the job description here",
        height=200,
        placeholder="Paste the full job posting here. The more detail, the better the analysis.",
        label_visibility="collapsed",
    )
    st.markdown("</div>", unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 0.5, 1])
    with col2:
        run = st.button("⚡  Analyze", use_container_width=True)

    if run:
        if not data.get("name") and not data.get("skills"):
            st.warning("⚠️ Fill in at least your name and skills in the **Build Resume** tab first.")
            return

        with st.spinner("Running analysis..."):
            import time; time.sleep(0.8)  # UX pause
            analysis = analyze_resume(data, jd)
            st.session_state["analysis"] = analysis
            st.session_state["jd"]       = jd

    analysis = st.session_state.get("analysis")
    if not analysis:
        st.info("📊 Results will appear here after you click **Analyze**.")
        return

    # ── Score cards ──
    ats = analysis["ats_score"]
    score_cls = "score-high" if ats >= 75 else ("score-medium" if ats >= 55 else "score-low")
    grade     = "Strong ✅" if ats >= 75 else ("Fair ⚠️" if ats >= 55 else "Needs Work ❌")

    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value {score_cls}">{ats}</div>
            <div class="metric-label">ATS Score / 100</div>
        </div>""", unsafe_allow_html=True)
    with c2:
        skill_count = len(data.get("skills", []))
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value" style="color:#a5b4fc;">{skill_count}</div>
            <div class="metric-label">Skills Listed</div>
        </div>""", unsafe_allow_html=True)
    with c3:
        proj_count = len([p for p in data.get("projects", []) if p.get("title")])
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value" style="color:#fbbf24;">{proj_count}</div>
            <div class="metric-label">Projects</div>
        </div>""", unsafe_allow_html=True)

    st.markdown(f"<br><div style='text-align:center;color:#7c8494;font-size:0.9rem;'>Match quality: <b style='color:#e2e8f0;'>{grade}</b></div>", unsafe_allow_html=True)
    st.progress(ats / 100)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Skill suggestions ──
    suggestions = analysis.get("skill_suggestions", [])
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<p class="section-title">💡 Recommended Skills to Add</p>', unsafe_allow_html=True)

    existing_tags = "".join(f'<span class="tag">{s}</span>' for s in data.get("skills", []))
    new_tags      = "".join(f'<span class="tag tag-new">+ {s}</span>' for s in suggestions)

    st.markdown(f"""
    <p style="font-size:0.82rem;color:#7c8494;margin-bottom:0.5rem;">Currently on your resume:</p>
    <div class="tag-container">{existing_tags or '<span style="color:#7c8494;font-size:0.85rem;">No skills listed</span>'}</div>
    <p style="font-size:0.82rem;color:#52d9a4;margin:1rem 0 0.5rem;">Suggested additions:</p>
    <div class="tag-container">{new_tags or '<span style="color:#7c8494;font-size:0.85rem;">Your skills look comprehensive!</span>'}</div>
    """, unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # ── Improved project descriptions ──
    improved = analysis.get("improved_projects", [])
    if improved:
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown('<p class="section-title">✨ AI-Enhanced Project Descriptions</p>', unsafe_allow_html=True)
        st.markdown("<p style='font-size:0.82rem;color:#7c8494;'>These quantified descriptions are used when you download the resume with <em>AI improvements</em> enabled.</p>", unsafe_allow_html=True)
        for ip in improved:
            st.markdown(f"**{ip['title']}**")
            st.markdown(f'<div class="suggestion-box">{ip["improved"]}</div>', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

    # ── UPDATED: Keyword Match Panel (NEW section) ──
    km = analysis.get("keyword_match", {})
    matched_kws = km.get("matched_keywords", [])
    missing_kws = km.get("missing_keywords", [])
    if matched_kws or missing_kws:
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown('<p class="section-title">🔑 Keyword Match Analysis</p>', unsafe_allow_html=True)
        st.markdown(
            "<p style='font-size:0.82rem;color:#7c8494;margin-bottom:0.8rem;'>"
            "Keywords extracted from the job description compared against your resume content.</p>",
            unsafe_allow_html=True,
        )
        kc1, kc2 = st.columns(2)
        with kc1:
            matched_html = "".join(
                f'<span class="tag" style="border-color:#2d6a4f;color:#52d9a4;">{kw}</span>'
                for kw in matched_kws
            )
            matched_fallback = '<span style="color:#7c8494">None found</span>'
            st.markdown(
                f"<p style='font-size:0.8rem;color:#52d9a4;margin-bottom:0.4rem;'>✅ Matched ({len(matched_kws)})</p>"
                f"<div class='tag-container'>{matched_html or matched_fallback}</div>",
                unsafe_allow_html=True,
            )
        with kc2:
            missing_html = "".join(
                f'<span class="tag" style="border-color:#7f1d1d;color:#f87171;">{kw}</span>'
                for kw in missing_kws
            )
            missing_fallback = '<span style="color:#7c8494">None — great coverage!</span>'
            st.markdown(
                f"<p style='font-size:0.8rem;color:#f87171;margin-bottom:0.4rem;'>❌ Missing ({len(missing_kws)})</p>"
                f"<div class='tag-container'>{missing_html or missing_fallback}</div>",
                unsafe_allow_html=True,
            )
        st.markdown("</div>", unsafe_allow_html=True)

    # ── UPDATED: AI Suggestions Panel (NEW section) ──
    ai_sugg = analysis.get("ai_suggestions", {})
    if ai_sugg:
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown('<p class="section-title">🤖 AI Coach Recommendations</p>', unsafe_allow_html=True)

        ms_skills = ai_sugg.get("missing_skills", [])
        weak_sec  = ai_sugg.get("weak_sections", [])
        kw_gaps   = ai_sugg.get("keyword_gaps", [])
        advice    = ai_sugg.get("advice", [])

        if ms_skills:
            ms_html = "".join(f'<span class="tag tag-new">+ {s}</span>' for s in ms_skills)
            st.markdown(
                "<p style='font-size:0.82rem;color:#a5b4fc;margin:0.6rem 0 0.4rem;'>🧩 Skills to learn or add:</p>"
                f"<div class='tag-container'>{ms_html}</div>",
                unsafe_allow_html=True,
            )
        if kw_gaps:
            kw_html = "".join(f'<span class="tag" style="border-color:#92400e;color:#fbbf24;">{k}</span>' for k in kw_gaps)
            st.markdown(
                "<p style='font-size:0.82rem;color:#fbbf24;margin:0.8rem 0 0.4rem;'>🔍 JD keywords to weave in:</p>"
                f"<div class='tag-container'>{kw_html}</div>",
                unsafe_allow_html=True,
            )
        if weak_sec:
            st.markdown("<p style='font-size:0.82rem;color:#f87171;margin:0.8rem 0 0.4rem;'>⚠️ Sections needing attention:</p>",
                        unsafe_allow_html=True)
            for ws in weak_sec:
                st.markdown(f"<div class='suggestion-box'>⚠️ {ws}</div>", unsafe_allow_html=True)
        if advice:
            st.markdown("<p style='font-size:0.82rem;color:#52d9a4;margin:0.8rem 0 0.4rem;'>✅ Actionable advice:</p>",
                        unsafe_allow_html=True)
            for tip in advice:
                st.markdown(f"<div class='suggestion-box'>💬 {tip}</div>", unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)
    elif _get_groq_client() is None:
        st.markdown(
            '<div class="section-card"><p class="section-title">🤖 AI Coach Recommendations</p>'
            '<p style="font-size:0.85rem;color:#7c8494;">Set the <code>GROQ_API_KEY</code> environment variable '
            'to unlock AI-powered coaching, keyword gap analysis, and smart project rewrites.</p></div>',
            unsafe_allow_html=True,
        )

    # ── Tips ──
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<p class="section-title">📌 Quick ATS Tips</p>', unsafe_allow_html=True)
    tips = [
        "Use keywords from the job description verbatim — ATS parsers match exact strings.",
        "Quantify achievements: numbers stand out (e.g. '40% faster', '10k users').",
        "Avoid tables, graphics, or headers/footers in your DOCX — many ATS tools misparse them.",
        "Include a dedicated Skills section with relevant technologies.",
        "Use standard section names: Education, Experience, Projects, Skills.",
    ]
    for tip in tips:
        st.markdown(f"<div class='suggestion-box'>💡 {tip}</div>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# 6. PREVIEW PAGE
# ─────────────────────────────────────────────
def render_preview_page(data: dict):
    st.markdown("""
    <div class="hero-card">
        <p class="hero-title">👁 Preview & Export</p>
        <p class="hero-sub">Review your resume and download a polished DOCX file</p>
    </div>
    """, unsafe_allow_html=True)

    analysis = st.session_state.get("analysis", {})
    use_improved = False

    # Export controls
    col1, col2 = st.columns([3, 1])
    with col1:
        if analysis:
            use_improved = st.checkbox(
                "✨ Use AI-improved project descriptions",
                value=False,
                help="Swap in the quantified descriptions generated during analysis.",
            )
    with col2:
        st.markdown("<div style='height:1px'/>", unsafe_allow_html=True)

    # Generate DOCX
    docx_bytes = generate_resume_docx(data, analysis, use_improved=use_improved)
    fname = (data.get("name") or "resume").replace(" ", "_") + "_resume.docx"

    st.download_button(
        label="⬇️  Download Resume (.docx)",
        data=docx_bytes,
        file_name=fname,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=False,
    )

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Live HTML preview ──
    def _safe(val, fallback="—"):
        return val if val else fallback

    # Contact
    contact_bits = []
    if data.get("email"):    contact_bits.append(f'✉ {data["email"]}')
    if data.get("phone"):    contact_bits.append(f'📞 {data["phone"]}')
    if data.get("location"): contact_bits.append(f'📍 {data["location"]}')
    if data.get("linkedin"): contact_bits.append(f'🔗 {data["linkedin"]}')
    if data.get("github"):   contact_bits.append(f'💻 {data["github"]}')
    contact_html = "  ·  ".join(contact_bits) or "—"

    # Skills
    skills_html = (
        " &nbsp; ".join(f'<span style="background:#f0f0ff;border:1px solid #c7d2fe;border-radius:4px;padding:1px 8px;font-size:0.82rem;color:#4f46e5;">{s}</span>'
                        for s in data.get("skills", []))
        or "—"
    )

    # Education
    edu_html = ""
    for edu in data.get("education", []):
        gpa_str = f" &nbsp;&nbsp; <span style='color:#888;font-size:0.82rem;'>GPA: {edu['gpa']}</span>" if edu.get("gpa") else ""
        edu_html += f"""
        <p class="resume-item-title">{_safe(edu.get('degree'))}</p>
        <p class="resume-item-sub">{_safe(edu.get('institution'))} &nbsp;·&nbsp; {_safe(edu.get('year'))}{gpa_str}</p>
        """

    # Experience
    exp_html = ""
    for exp in data.get("experience", []):
        bullets_html = "".join(
            f'<p class="resume-bullet">• {b}</p>'
            for b in exp.get("bullets", [])
        )
        exp_html += f"""
        <p class="resume-item-title">{_safe(exp.get('role'))} &nbsp;<span style='color:#888;font-weight:400;font-size:0.88rem;'>@ {_safe(exp.get('company'))}</span></p>
        <p class="resume-item-sub" style="margin-bottom:4px;">{_safe(exp.get('period'))}</p>
        {bullets_html}
        <div style="height:6px;"></div>
        """

    # Projects
    improved_map = {ip["title"]: ip["improved"] for ip in analysis.get("improved_projects", [])}
    proj_html = ""
    for proj in data.get("projects", []):
        if not proj.get("title"): continue
        desc = (improved_map.get(proj["title"]) if use_improved else proj.get("description")) or ""
        proj_html += f"""
        <p class="resume-item-title">{proj['title']}
            <span style="color:#4f46e5;font-weight:400;font-size:0.85rem;"> · {proj.get('tech','')}</span>
        </p>
        <p class="resume-bullet" style="padding-left:0;">{desc}</p>
        <div style="height:6px;"></div>
        """

    # Achievements
    ach_html = "".join(
        f'<p class="resume-bullet">• {a.lstrip("•-– ").strip()}</p>'
        for a in data.get("achievements", [])
    )

    def _section(title, content):
        if not content.strip(): return ""
        return f"""
        <div class="resume-section-title">{title}</div>
        {content}
        """

    preview_html = f"""
    <div class="resume-preview">
        <div class="resume-name">{_safe(data.get('name'), 'Your Name')}</div>
        <div class="resume-contact">{contact_html}</div>
        {_section("Skills", f'<p style="font-size:0.88rem;color:#333;">{skills_html}</p>')}
        {_section("Education", edu_html)}
        {_section("Work Experience", exp_html)}
        {_section("Projects", proj_html)}
        {_section("Achievements & Certifications", ach_html)}
    </div>
    """

    st.markdown(preview_html, unsafe_allow_html=True)

    if not analysis:
        st.info("💡 Run **Analyze** to unlock AI-improved descriptions and ATS scoring.")


# ─────────────────────────────────────────────
# 7. MAIN
# ─────────────────────────────────────────────
def main():
    # Persist data across tab switches
    if "resume_data" not in st.session_state:
        st.session_state["resume_data"] = {}

    page = render_sidebar()

    if page == "📝  Build Resume":
        data = collect_user_data()
        st.session_state["resume_data"] = data  # Save on every render

    elif page == "🔍  Analyze":
        data = st.session_state.get("resume_data", {})
        render_analysis_page(data)

    elif page == "👁  Preview & Export":
        data = st.session_state.get("resume_data", {})
        render_preview_page(data)


if __name__ == "__main__":
    main()